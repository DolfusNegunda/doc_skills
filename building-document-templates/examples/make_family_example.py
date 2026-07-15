"""Generate a neutral, synthetic FAMILY example source document (no client data).

Writes `family_source.docx` — a fictional "Project Status Report" for the made-up
company *Northwind Traders* (project PRJ-0042). It has the shape a real family
exemplar has: a data-bound cover (docProps title/subject), a prose summary, and
three variable-count tables (team, deliverables, metrics) — exactly the things the
family model turns into cover-property fields, text fields, and ROW-GROUPS.

It is reproducible (no binary committed to git) and carries only invented data, so
the derived template and the worked run in README.md are safe to publish. Run it,
then follow README.md to `propose -> build --family -> fill -> validate`.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_table(doc, header, rows):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        t.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    return t


def main():
    doc = Document()

    # Data-bound cover: these render on a Word cover page from docProps, NOT body runs.
    doc.core_properties.title = "Northwind Traders — Widget Line Status"
    doc.core_properties.subject = "Monthly Project Status Report"

    title = doc.add_heading("Northwind Traders — Widget Line Status", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Monthly Project Status Report — PRJ-0042")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Report date: 2024/01/15").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(
        "The Widget Assembly Line automation reached 70% of planned throughput this "
        "period. Integration testing for the PRJ-0042 conveyor module is complete and "
        "the Northwind Traders operations team has begun acceptance trials.")

    doc.add_heading("Team", level=1)
    add_table(doc, ["Name", "Role", "Allocation"], [
        ["Ada Lovelace", "Project Lead", "0.5 FTE"],
        ["Grace Hopper", "Automation Engineer", "1.0 FTE"],
    ])

    doc.add_heading("Deliverables", level=1)
    add_table(doc, ["Item", "Description", "Status"], [
        ["Conveyor module", "Integrate the PRJ-0042 conveyor into the line", "Done"],
        ["Control software", "Widget Assembly Line PLC program", "In progress"],
        ["Acceptance report", "Sign-off from Northwind Traders operations", "Not started"],
    ])

    doc.add_heading("Key Metrics", level=1)
    add_table(doc, ["Metric", "Value"], [
        ["Throughput vs. plan", "70%"],
        ["Open defects", "12"],
    ])

    out = Path(__file__).with_name("family_source.docx")
    doc.save(str(out))
    print(f"wrote {out}")


if __name__ == "__main__":
    main()
