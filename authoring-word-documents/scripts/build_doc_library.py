"""Generate the built-in DOCX template library — branded documents a small model FILLS.

The observed small-model failure is a structurally-sound document on the stock
python-docx theme: default Office palette, no logo, empty footer, empty metadata —
"untouched Word 2010", not corporate output. These templates fix that: every one has
a real style architecture (Title/Subtitle/Heading 1-3/List Bullet/Caption), brand
fonts and colors, a logo header, a legal footer with page numbers, filled document
properties, and `{{ tagged }}` variable text with row-group tables where content
repeats. Each is registered with a hand-authored manifest in the same run, so
template and manifest can never drift. Fill with the standard engine:

    python ../building-document-templates/scripts/registry.py show --client _builtin --doc-type business_report
    python ../building-document-templates/scripts/registry.py scaffold --builtin business_report --out content.json
    python ../building-document-templates/scripts/fill.py --client _builtin --doc-type business_report \
        --data content.json --out out.docx
    python ../building-document-templates/scripts/validate.py out.docx --template <registry>/_builtin/business_report/template.docx

Branding is data, not code — see ../../brands/README.md.

    python scripts/build_doc_library.py                     # all templates, default brand
    python scripts/build_doc_library.py --brand path/to/client-pack   # re-skin for a client
    python scripts/build_doc_library.py --only memo         # one template

Templates: business_report, memo, meeting_minutes, one_pager. Requires python-docx.
"""
from __future__ import annotations

import argparse
import io
import json
import os
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

SKILL_ROOT = Path(__file__).resolve().parent.parent
REPO_ROOT = SKILL_ROOT.parent
BRANDS = REPO_ROOT / "brands"
DEFAULT_REGISTRY = REPO_ROOT / "building-document-templates" / "registry"


def tag(name: str) -> str:
    return "{{ " + name + " }}"


# ---------------- Brand pack (same contract as the other builders) ----------------

def _deep_merge(base: dict, over: dict) -> dict:
    out = dict(base)
    for k, v in over.items():
        out[k] = _deep_merge(base[k], v) if isinstance(v, dict) and isinstance(base.get(k), dict) else v
    return out


def load_brand(name_or_path: str | None) -> dict:
    default = json.loads((BRANDS / "default" / "brand.json").read_text(encoding="utf-8"))
    default["_dir"] = BRANDS / "default"
    if not name_or_path or name_or_path == "default":
        return default
    p = Path(name_or_path)
    if p.suffix == ".json" and p.exists():
        pack_file, pack_dir = p, p.parent
    elif p.is_dir() and (p / "brand.json").exists():
        pack_file, pack_dir = p / "brand.json", p
    elif (BRANDS / name_or_path / "brand.json").exists():
        pack_dir = BRANDS / name_or_path
        pack_file = pack_dir / "brand.json"
    else:
        known = sorted(d.name for d in BRANDS.iterdir() if (d / "brand.json").exists())
        sys.exit(f"Unknown brand '{name_or_path}'. Available: {', '.join(known)}")
    pack = _deep_merge(default, json.loads(pack_file.read_text(encoding="utf-8")))
    pack["_dir"] = pack_dir
    return pack


def rgb(hex_str: str) -> RGBColor:
    return RGBColor.from_string(hex_str.lstrip("#"))


class Style:
    def __init__(self, brand: dict):
        c = brand["colors"]
        self.ink = rgb(c["ink"])
        self.muted = rgb(c["muted"])
        self.accent = rgb(c["primary"])
        self.dark = rgb(c["dark"])
        self.panel_hex = c["panel"].lstrip("#")
        self.dark_hex = c["dark"].lstrip("#")
        self.font = brand["fonts"]["body"]
        self.font_h = brand["fonts"]["heading"]
        year = date.today().year
        f = brand.get("footer", {})
        cop = (f.get("copyright") or "").replace("{year}", str(year)) \
                                        .replace("{company}", brand.get("display_name", ""))
        self.copyright = cop
        self.confidentiality = f.get("confidentiality", "")
        self.company = brand.get("display_name", "")
        logo_rel = brand.get("logo")
        self.logo_bytes = (brand["_dir"] / logo_rel).read_bytes() if logo_rel else None


# ---------------- Document scaffolding ----------------

def _set_font(style, name, size=None, color=None, bold=None):
    style.font.name = name
    # cover complex-script/eastasia so the font sticks everywhere
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)
    if size is not None:
        style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold


def configure_styles(doc: Document, st: Style) -> None:
    """Real style architecture: everything downstream uses named styles, never
    direct formatting — so a filled document restyles coherently."""
    s = doc.styles
    _set_font(s["Normal"], st.font, size=10.5, color=st.ink)
    s["Normal"].paragraph_format.space_after = Pt(8)
    s["Normal"].paragraph_format.line_spacing = 1.18

    _set_font(s["Title"], st.font_h, size=28, color=st.dark, bold=True)
    s["Title"].paragraph_format.space_before = Pt(6)
    s["Title"].paragraph_format.space_after = Pt(4)

    _set_font(s["Subtitle"], st.font, size=13, color=st.muted)
    s["Subtitle"].paragraph_format.space_after = Pt(18)

    for name, size, color, before, after in (
            ("Heading 1", 15, st.dark, 18, 6),
            ("Heading 2", 12.5, st.accent, 14, 4),
            ("Heading 3", 11, st.ink, 10, 3)):
        _set_font(s[name], st.font_h, size=size, color=color, bold=True)
        s[name].paragraph_format.space_before = Pt(before)
        s[name].paragraph_format.space_after = Pt(after)

    _set_font(s["List Bullet"], st.font, size=10.5, color=st.ink)
    s["List Bullet"].paragraph_format.space_after = Pt(4)
    _set_font(s["Caption"], st.font, size=9, color=st.muted)


def _page_number_run(paragraph):
    run = paragraph.add_run()
    for el, text in (("w:fldChar", None), ("w:instrText", "PAGE"), ("w:fldChar", None)):
        e = OxmlElement(el)
        if el == "w:fldChar":
            e.set(qn("w:fldCharType"), "begin" if text is None and len(run._r) == 0 else "end")
        else:
            e.set(qn("xml:space"), "preserve")
            e.text = " PAGE "
        run._r.append(e)
    return run


def setup_page_and_chrome(doc: Document, st: Style) -> None:
    """A4, sensible margins, logo header, legal footer with page number."""
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(2.2)
    sec.top_margin, sec.bottom_margin = Cm(1.8), Cm(1.8)

    if st.logo_bytes:
        hp = sec.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.add_run().add_picture(io.BytesIO(st.logo_bytes), height=Cm(0.9))

    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    legal = "   ·   ".join(x for x in (st.copyright, st.confidentiality) if x)
    r = fp.add_run((legal + "   ·   " if legal else "") + "Page ")
    r.font.size = Pt(7.5)
    r.font.color.rgb = st.muted
    r.font.name = st.font
    pr = _page_number_run(fp)
    pr.font.size = Pt(7.5)
    pr.font.color.rgb = st.muted
    pr.font.name = st.font


def _shade(cell, hex_fill: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shd)


def meta_table(doc: Document, st: Style, pairs: list[tuple[str, str]]):
    """Label/value block under the title — the document's identity card."""
    t = doc.add_table(rows=len(pairs), cols=2)
    t.style = doc.styles["Table Grid"]
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.columns[0].width = Cm(4.2)
    t.columns[1].width = Cm(12.4)
    for row, (label, value) in zip(t.rows, pairs):
        row.cells[0].width, row.cells[1].width = Cm(4.2), Cm(12.4)
        _shade(row.cells[0], st.panel_hex)
        p0 = row.cells[0].paragraphs[0]
        r0 = p0.add_run(label)
        r0.font.bold = True
        r0.font.size = Pt(9.5)
        r0.font.color.rgb = st.dark
        p1 = row.cells[1].paragraphs[0]
        r1 = p1.add_run(value)
        r1.font.size = Pt(9.5)
    return t


def data_table(doc: Document, st: Style, columns: list[str], row_tags: list[str]):
    """Header row + ONE template row whose cells hold the row-group tags —
    fill.py clones the row per data item (real repeating rows)."""
    t = doc.add_table(rows=2, cols=len(columns))
    t.style = doc.styles["Table Grid"]
    for cell, col in zip(t.rows[0].cells, columns):
        _shade(cell, st.dark_hex)
        r = cell.paragraphs[0].add_run(col)
        r.font.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = rgb("#FFFFFF")
    for cell, rt in zip(t.rows[1].cells, row_tags):
        r = cell.paragraphs[0].add_run(rt)
        r.font.size = Pt(9.5)
    return t


def add_body(doc, text, style=None):
    p = doc.add_paragraph(text)
    if style:
        p.style = doc.styles[style]
    return p


# ---------------- Field helper ----------------

def F(name, example, guidance, *, type="text", required=True):
    return {"name": name, "type": type, "example": example, "guidance": guidance,
            "required": required, "media_part": ""}


SOURCE_TERMS = ["Acme Mining", "Jane Mokoena", "Sipho Dlamini"]


# ---------------- Template definitions ----------------

def build_business_report(doc, st):
    doc.add_paragraph(tag("report_title"), style="Title")
    doc.add_paragraph(tag("report_subtitle"), style="Subtitle")
    meta_table(doc, st, [
        ("Prepared for", tag("client_name")),
        ("Project", tag("project_name")),
        ("Project code", tag("project_code")),
        ("Date", tag("report_date")),
        ("Author", tag("author_line")),
        ("Status", tag("report_status")),
    ])
    doc.add_paragraph("Executive summary", style="Heading 1")
    add_body(doc, tag("executive_summary"))
    doc.add_paragraph("Background", style="Heading 1")
    add_body(doc, tag("background"))
    doc.add_paragraph("Key findings", style="Heading 1")
    doc.add_paragraph(tag("findings"), style="List Bullet")
    doc.add_paragraph("Analysis", style="Heading 1")
    add_body(doc, tag("analysis_detail"))
    doc.add_paragraph("Deliverables", style="Heading 1")
    data_table(doc, st, ["Deliverable", "Description", "Status"],
               [tag("deliverable_item"), tag("deliverable_description"), tag("deliverable_status")])
    doc.add_paragraph("Recommendations", style="Heading 1")
    doc.add_paragraph(tag("recommendations"), style="List Bullet")
    doc.add_paragraph("Next steps", style="Heading 1")
    doc.add_paragraph(tag("next_steps"), style="List Bullet")

    fields = [
        F("report_title", "Q3 2026 Operations Review", "Document title, under ~10 words."),
        F("report_subtitle", "Performance, drivers, and the decisions required for Q4.", "One-sentence framing under the title."),
        F("client_name", "Acme Mining", "Client or audience organisation."),
        F("project_name", "Fleet Optimisation Study", "Engagement/project name."),
        F("project_code", "RI.1234", "Internal project code."),
        F("report_date", "15 October 2026", "Issue date."),
        F("author_line", "Jane Mokoena, Engagement Lead", "Author name and role."),
        F("report_status", "Final for review", "Draft/Final + review state."),
        F("executive_summary", "The quarter closed 4% above plan on the strength of two enterprise renewals. Margin improved 1.2 points; the one deteriorating indicator is customer NPS, for which a funded recovery plan is underway.", "3-5 sentences: findings first, then implication, then the ask."),
        F("background", "This review covers the third quarter of 2026 and follows the methodology agreed in the January planning cycle.", "1-2 short paragraphs of context: scope, period, method."),
        F("findings", "Cost per tonne fell 14% in the phased scenario", "3-6 finding bullets — assertions, not topics.", type="list"),
        F("analysis_detail", "The cost improvement is driven primarily by cycle-time consistency rather than fleet size: variance between shifts fell 40% after dispatch rule changes.", "The analytical detail behind the findings; several sentences."),
        F("recommendations", "Commit to the phased transition starting north pit Q2 2027", "3-5 recommendation bullets, ranked.", type="list"),
        F("next_steps", "Board decision on phase 1 capex — Sipho Dlamini — by 31 January", "Action bullets with owners and dates.", type="list"),
    ]
    row_groups = [{"name": "deliverables",
                   "columns": ["deliverable_item", "deliverable_description", "deliverable_status"]}]
    fields += [
        F("deliverable_item", "Simulation model", "Row group 'deliverables': short deliverable name.", required=False),
        F("deliverable_description", "Validated model of current haulage operations", "What the deliverable is.", required=False),
        F("deliverable_status", "Delivered", "Delivered / In progress / At risk.", required=False),
    ]
    return fields, row_groups, "Formal business/project report: exec summary, background, findings, analysis, deliverables table, recommendations."


def build_memo(doc, st):
    doc.add_paragraph(tag("memo_title"), style="Title")
    doc.add_paragraph("Internal memorandum", style="Subtitle")
    meta_table(doc, st, [
        ("To", tag("to_line")),
        ("From", tag("from_line")),
        ("Date", tag("memo_date")),
        ("Re", tag("subject_line")),
    ])
    doc.add_paragraph("Purpose", style="Heading 1")
    add_body(doc, tag("purpose"))
    doc.add_paragraph("Background", style="Heading 1")
    add_body(doc, tag("background"))
    doc.add_paragraph("Key points", style="Heading 1")
    doc.add_paragraph(tag("points"), style="List Bullet")
    doc.add_paragraph("Action requested", style="Heading 1")
    add_body(doc, tag("action_requested"))

    fields = [
        F("memo_title", "Change to the Q4 delivery schedule", "Memo title — the decision or topic."),
        F("to_line", "Executive Committee", "Recipient person/group."),
        F("from_line", "Jane Mokoena, Engagement Lead", "Sender and role."),
        F("memo_date", "3 November 2026", "Memo date."),
        F("subject_line", "Resourcing decision required before 15 November", "One-line subject."),
        F("purpose", "This memo requests approval to bring forward two delivery hires to protect the Q4 start dates.", "1-2 sentences: why this memo exists."),
        F("background", "Two client starts moved into Q4, concentrating delivery demand in an eight-week window.", "Short context paragraph."),
        F("points", "Current bench covers only one of the three concurrent starts", "3-6 bullets carrying the argument.", type="list"),
        F("action_requested", "Approve the two hires by 15 November so notice periods land before the January starts.", "The specific ask, with the deadline."),
    ]
    return fields, [], "Internal memo: to/from block, purpose, background, key points, action requested."


def build_meeting_minutes(doc, st):
    doc.add_paragraph(tag("meeting_title"), style="Title")
    doc.add_paragraph("Minutes of meeting", style="Subtitle")
    meta_table(doc, st, [
        ("Date & time", tag("meeting_datetime")),
        ("Location", tag("location")),
        ("Chair", tag("chair")),
        ("Minute-taker", tag("minute_taker")),
    ])
    doc.add_paragraph("Attendees", style="Heading 1")
    data_table(doc, st, ["Name", "Role / Organisation"],
               [tag("attendee_name"), tag("attendee_role")])
    doc.add_paragraph("Agenda", style="Heading 1")
    doc.add_paragraph(tag("agenda_items"), style="List Bullet")
    doc.add_paragraph("Discussion", style="Heading 1")
    add_body(doc, tag("discussion_summary"))
    doc.add_paragraph("Decisions", style="Heading 1")
    doc.add_paragraph(tag("decisions"), style="List Bullet")
    doc.add_paragraph("Actions", style="Heading 1")
    data_table(doc, st, ["Action", "Owner", "Due"],
               [tag("action_item"), tag("action_owner"), tag("action_due")])

    fields = [
        F("meeting_title", "Fleet Optimisation Study — Steering Committee #4", "Meeting name and number."),
        F("meeting_datetime", "3 November 2026, 10:00–11:00 SAST", "Date and time with timezone."),
        F("location", "Client boardroom / Teams", "Where the meeting happened."),
        F("chair", "Sipho Dlamini", "Who chaired."),
        F("minute_taker", "Jane Mokoena", "Who took minutes."),
        F("agenda_items", "Progress against sprint 2 plan", "The agenda, one bullet per item.", type="list"),
        F("discussion_summary", "The committee reviewed sprint-2 progress; model validation is a week ahead of plan while data access for the north pit remains outstanding.", "Narrative summary of the discussion; a few sentences per topic."),
        F("decisions", "Approved the revised scenario list (five scenarios)", "One bullet per decision taken.", type="list"),
    ]
    row_groups = [
        {"name": "attendees", "columns": ["attendee_name", "attendee_role"]},
        {"name": "actions", "columns": ["action_item", "action_owner", "action_due"]},
    ]
    fields += [
        F("attendee_name", "Jane Mokoena", "Row group 'attendees': person's name.", required=False),
        F("attendee_role", "Engagement Lead, consultant", "Role and organisation.", required=False),
        F("action_item", "Chase north-pit data extract", "Row group 'actions': the action.", required=False),
        F("action_owner", "Sipho Dlamini", "Who owns it.", required=False),
        F("action_due", "10 November 2026", "Due date.", required=False),
    ]
    return fields, row_groups, "Meeting minutes: attendees table, agenda, discussion, decisions, actions table."


def build_one_pager(doc, st):
    doc.add_paragraph(tag("title"), style="Title")
    doc.add_paragraph(tag("subtitle"), style="Subtitle")
    doc.add_paragraph("The ask", style="Heading 1")
    add_body(doc, tag("the_ask"))
    doc.add_paragraph("Why now", style="Heading 1")
    doc.add_paragraph(tag("why_now"), style="List Bullet")
    doc.add_paragraph("What it takes", style="Heading 1")
    add_body(doc, tag("what_it_takes"))
    doc.add_paragraph("Expected impact", style="Heading 1")
    doc.add_paragraph(tag("impact_points"), style="List Bullet")
    doc.add_paragraph("Contact", style="Heading 1")
    add_body(doc, tag("contact_line"))

    fields = [
        F("title", "Trolley-Assist Transition — Decision Brief", "The one_pager's title."),
        F("subtitle", "One page on the decision, the cost, and the payoff.", "One-line subtitle."),
        F("the_ask", "Approve phase-1 capex of $2.4m to start the north-pit transition in Q2 2027.", "The single decision requested, 1-2 sentences."),
        F("why_now", "Diesel exposure grows every quarter the decision waits", "3-4 urgency bullets.", type="list"),
        F("what_it_takes", "Eight weeks of engineering design, one substation upgrade, and a phased fleet retrofit across 2027.", "Resources/steps in 2-3 sentences."),
        F("impact_points", "14% lower cost per tonne at current diesel prices", "3-4 quantified impact bullets.", type="list"),
        F("contact_line", "Jane Mokoena · jane.mokoena@example.com · +27 11 000 0000", "Who to talk to."),
    ]
    return fields, [], "Decision one-pager: the ask, why now, what it takes, impact, contact."


TEMPLATES = {
    "business_report": build_business_report,
    "memo": build_memo,
    "meeting_minutes": build_meeting_minutes,
    "one_pager": build_one_pager,
}


# ---------------- Build + register ----------------

def build_one(name: str, brand: dict, registry: Path, owner: str, created: str) -> Path:
    st = Style(brand)
    doc = Document()
    configure_styles(doc, st)
    setup_page_and_chrome(doc, st)
    fields, row_groups, description = TEMPLATES[name](doc, st)

    core = doc.core_properties
    core.title = f"{name} template ({brand['name']} brand)"
    core.author = st.company or "document-template-suite"
    core.comments = "Built-in template from the document-template suite. Fill via fill.py; do not edit placeholders by hand."

    dest = registry / "_builtin" / name
    dest.mkdir(parents=True, exist_ok=True)
    tpl = dest / "template.docx"
    doc.save(str(tpl))

    manifest = {
        "template_id": f"_builtin/{name}",
        "client": "_builtin",
        "doc_type": name,
        "format": "docx",
        "template_file": "template.docx",
        "source_file": f"generated by authoring-word-documents/scripts/build_doc_library.py (brand '{brand['name']}')",
        "version": "1.0.0",
        "owner": owner,
        "created": created,
        "changelog": [f"{created}: generated with brand pack '{brand['name']}'"],
        "description": description,
        "fields": fields,
        "row_groups": row_groups,
        "source_terms": SOURCE_TERMS,
    }
    (dest / "manifest.json").write_text(json.dumps(manifest, indent=2, ensure_ascii=False),
                                        encoding="utf-8")
    print(f"Registered _builtin/{name}: {len(fields)} fields, {len(row_groups)} row group(s) -> {dest}")
    return dest


def main():
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--brand", default="default", help="Brand pack name under brands/ (or a path)")
    ap.add_argument("--only", choices=sorted(TEMPLATES), help="Build a single template")
    ap.add_argument("--registry", default=os.environ.get("TEMPLATE_REGISTRY", str(DEFAULT_REGISTRY)),
                    help="Registry root (default: building-document-templates/registry or $TEMPLATE_REGISTRY)")
    ap.add_argument("--owner", default="", help="Recorded in each manifest")
    ap.add_argument("--created", default=date.today().isoformat())
    args = ap.parse_args()

    brand = load_brand(args.brand)
    registry = Path(args.registry)
    for name in ([args.only] if args.only else sorted(TEMPLATES)):
        build_one(name, brand, registry, args.owner, args.created)
    print("\nDiscover:  python ../building-document-templates/scripts/registry.py list")


if __name__ == "__main__":
    main()
